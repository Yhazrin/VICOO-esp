# -*- coding: utf-8 -*-
"""Generate bilingual video script docx with admin section."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT_PATH = r"c:\Users\12065\PycharmProjects\VICOO-esp\期末pre+视频文字稿（含管理后台-中英对照）.docx"

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_title(text, size=16, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_section_header(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_lang_label(label):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_body(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_divider():
    p = doc.add_paragraph("—" * 40)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        run.font.size = Pt(9)


add_title("VICOO 期末 Demo 视频文字稿")
add_title("（含管理后台 · 5 分钟 · 中英对照修订版）", size=12)
doc.add_paragraph()

sections = [
    {
        "time": "0:00–0:20  Opening / 开场",
        "en": [
            "Welcome to the VICOO system demonstration.",
            "This video shows how VICOO helps users discover sustainable fashion products, verify supply-chain information, and connect purchasing decisions with environmental impact.",
            "The demonstration follows a typical journey, from product browsing to traceability exploration, admin operations, and platform features.",
        ],
        "zh": [
            "欢迎观看 VICOO 系统演示。",
            "本视频将展示 VICOO 如何帮助用户发现可持续时尚产品、验证供应链信息，并将消费决策与环境影响联系起来。",
            "整个演示将按照典型用户流程展开，从产品浏览、traceability exploration、后台运营到平台功能展示。",
        ],
    },
    {
        "time": "0:20–0:42  Homepage / 首页",
        "en": [
            "We start from the VICOO homepage.",
            "The homepage introduces the platform's visual identity and guides users toward sustainable fashion discovery.",
            "From here, users can enter the shopping experience and explore impact-related products.",
        ],
        "zh": [
            "首先我们进入 VICOO 首页。",
            "首页展示平台整体视觉风格，并引导用户进入可持续时尚探索体验。",
            "用户可以从这里进入购物流程，并继续探索 impact-related products。",
        ],
    },
    {
        "time": "0:42–1:12  Impact Shop",
        "en": [
            "Next, we enter the Impact Shop.",
            "Users can browse products connected with sustainability or charity narratives. Product cards display images, descriptions, prices, and impact-related labels.",
            "Now we open one impact product and its detail page.",
        ],
        "zh": [
            "接下来我们进入 Impact Shop。",
            "用户可以浏览与 sustainability 或 charity narratives 相关的产品；产品卡片展示图片、简介、价格及 impact-related labels。",
            "现在我们打开一个 impact product 的详情页。",
        ],
    },
    {
        "time": "1:12–1:48  Product Detail Page / 产品详情页",
        "en": [
            "On the product detail page, users see images, sustainability information, and impact statements.",
            "Material sources, logistics stages, and certification records help users compare sustainability claims.",
        ],
        "zh": [
            "详情页展示图片、sustainability information 与 impact statements。",
            "材料来源、物流阶段与认证记录帮助用户理解并比较 sustainability claims。",
        ],
    },
    {
        "time": "1:48–2:38  Traceability Globe and Timeline / 追溯 Globe 与时间线",
        "en": [
            "Now we move to the traceability section.",
            "The Traceability Globe visualizes supply-chain stages as interactive geographic nodes—such as material sourcing, manufacturing, or logistics.",
            "When users select a stage, related timeline information appears below, helping users understand the supply chain as a connected journey.",
            "This is one of VICOO's core features: transforming structured supply-chain data into a clear visual explanation.",
        ],
        "zh": [
            "现在我们进入 traceability section。",
            "Traceability Globe 将供应链阶段可视化为交互式地理节点，例如 material sourcing、manufacturing 或 logistics。",
            "用户选择某一阶段时，下方同步显示 timeline information，帮助理解供应链的连续旅程。",
            "这是 VICOO 的核心功能之一：将结构化供应链数据转化为清晰直观的可视化说明。",
        ],
    },
    {
        "time": "2:38–3:03  AI Assistant / AI 助手",
        "en": [
            "Next, we demonstrate the AI Assistant.",
            "Users can ask product or sustainability questions in natural language, such as product materials or traceability information.",
            "The assistant works as a support layer over structured product information.",
        ],
        "zh": [
            "接下来我们展示 AI Assistant。",
            "用户可通过自然语言询问产品或 sustainability 相关问题，例如产品材料或 traceability information。",
            "AI Assistant 作为结构化产品信息的辅助层运行。",
        ],
    },
    {
        "time": "3:03–3:22  Donation and Charity Interaction / 公益与捐赠互动",
        "en": [
            "VICOO also includes charity-oriented interaction.",
            "Users can view donation-related content and understand how products or campaigns connect with social contribution.",
        ],
        "zh": [
            "VICOO 还包含 charity-oriented interaction。",
            "用户可以查看 donation-related content，并了解产品或 campaign 如何与 social contribution 建立联系。",
        ],
    },
    {
        "time": "3:22–3:47  Admin Dashboard / 管理后台（新增）",
        "en": [
            "Beyond the consumer site, VICOO provides an admin dashboard at /admin/ for editors and administrators.",
            "The dashboard summarizes platform activity. Staff can manage products, orders, campaigns, and donations from dedicated modules.",
            "For circular commerce, clothing intake requests can be reviewed and linked to listed products, closing the donate-to-resale loop.",
            "Role-based access and audit logs help keep operational changes traceable.",
        ],
        "zh": [
            "除用户端外，VICOO 还提供 /admin/ 管理后台，供编辑与管理员使用。",
            "Dashboard 汇总平台概况；运营人员可在独立模块中管理商品、订单、活动与捐赠。",
            "在循环商业场景中，衣物回收受理可审核并关联上架商品，形成「捐赠—再售」闭环。",
            "基于角色的权限与操作审计，保证后台变更可追溯。",
        ],
    },
    {
        "time": "3:47–4:07  System Reliability and Deployment / 系统部署与稳定性",
        "en": [
            "VICOO runs on FastAPI, MySQL, Redis, and Docker Compose.",
            "The same backend APIs serve both the public site and the admin dashboard.",
        ],
        "zh": [
            "系统基于 FastAPI、MySQL、Redis 与 Docker Compose 部署。",
            "同一套后端 API 同时服务用户端与管理后台。",
        ],
    },
    {
        "time": "4:07–4:30  Closing / 结尾",
        "en": [
            "To conclude, VICOO combines product discovery, supply-chain traceability, charity participation, AI guidance, and admin operations in one platform.",
            "Thank you for watching our demonstration.",
        ],
        "zh": [
            "总体而言，VICOO 将产品发现、供应链追溯、公益参与、AI 辅助与后台运营整合在同一平台中。",
            "感谢观看我们的演示。",
        ],
    },
]

for sec in sections:
    add_section_header(sec["time"])
    add_lang_label("English")
    for line in sec["en"]:
        add_body(line)
    add_lang_label("中文")
    for line in sec["zh"]:
        add_body(line)
    add_divider()

add_section_header("录屏提示 / Recording Notes")
add_lang_label("English")
add_body(
    "Total target duration: ~4:30–4:50 (within 5 minutes). "
    "Admin segment: login → Dashboard → Products/Orders → Clothing Donations or Audit Log (~6s each screen)."
)
add_lang_label("中文")
add_body(
    "总时长目标：约 4:30–4:50（控制在 5 分钟内）。"
    "管理后台段：登录 → Dashboard → 商品/订单 → 衣物回收或审计日志（每屏约 6 秒）。"
)

doc.save(OUT_PATH)
print("saved:", OUT_PATH)
